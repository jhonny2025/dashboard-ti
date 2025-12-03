from docx import Document

def generar_docx(data, ruta):
    doc = Document()

    doc.add_heading('PROYECTO', level=1)

    for clave, valor in data.items():
        doc.add_heading(clave.upper(), level=2)
        doc.add_paragraph(str(valor))

    doc.save(ruta)
